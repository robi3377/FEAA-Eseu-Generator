import io
import os
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fpdf import FPDF


def _get_font_dir() -> str:
    if sys.platform == "win32":
        return os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
    return "/usr/share/fonts/truetype/dejavu"


def _replace_latin_a_with_cyrillic(text: str) -> str:
    """Replace all Latin 'a' with Cyrillic 'а' (U+0430) and 'A' with 'А' (U+0410)."""
    text = text.replace("a", "\u0430")
    text = text.replace("A", "\u0410")
    return text


def _is_title(stripped: str) -> bool:
    if len(stripped) >= 100 or stripped.endswith(".") or stripped.endswith(","):
        return False
    if len(stripped.split()) > 10:
        return False
    return (
        stripped.startswith("Introducere")
        or stripped.startswith("Concluzii")
        or stripped.startswith("Bibliografie")
        or stripped.startswith("Secțiune")
        or stripped.startswith("Secţiune")
        or (len(stripped.split()) <= 8 and len(stripped) > 0 and stripped[0].isupper())
    )


def generate_docx(text: str, apply_a_replacement: bool = True) -> bytes:
    """Generate a .docx file from text."""
    doc = Document()

    # Remove any document protection / read-only flags
    settings = doc.settings.element
    for tag in settings.findall(qn('w:documentProtection')):
        settings.remove(tag)
    # Mark document as final=false so Word doesn't open read-only
    doc.core_properties.content_status = ""

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(2)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    if apply_a_replacement:
        text = _replace_latin_a_with_cyrillic(text)

    lines = text.split("\n")
    first_content_line = True

    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            first_content_line = False
            continue

        para = doc.add_paragraph()

        # First non-empty line is the project title
        if first_content_line:
            first_content_line = False
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(12)
        elif _is_title(stripped):
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(14)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
        else:
            run = para.add_run(stripped)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_pdf_from_text(text: str, title: str = "Document", apply_a_replacement: bool = False) -> bytes:
    """Generate a PDF from text using fpdf2 (pure Python, no system deps)."""
    if apply_a_replacement:
        text = _replace_latin_a_with_cyrillic(text)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(20, 20, 20)
    pdf.add_page()

    # Register a Unicode TTF font
    font_dir = _get_font_dir()
    if sys.platform == "win32":
        pdf.add_font("TNR", "", os.path.join(font_dir, "times.ttf"))
        pdf.add_font("TNR", "B", os.path.join(font_dir, "timesbd.ttf"))
        font_name = "TNR"
    else:
        pdf.add_font("DejaVu", "", os.path.join(font_dir, "DejaVuSans.ttf"))
        pdf.add_font("DejaVu", "B", os.path.join(font_dir, "DejaVuSans-Bold.ttf"))
        font_name = "DejaVu"

    lines = text.split("\n")
    first_content_line = True

    for line in lines:
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            first_content_line = False
            continue

        # First non-empty line is the project title
        if first_content_line:
            first_content_line = False
            pdf.set_font(font_name, "B", 16)
            pdf.multi_cell(0, 10, stripped, align="C")
            pdf.ln(6)
        elif _is_title(stripped):
            pdf.ln(5)
            pdf.set_font(font_name, "B", 14)
            pdf.multi_cell(0, 7, stripped, align="C")
            pdf.ln(3)
        else:
            pdf.set_font(font_name, "", 12)
            pdf.multi_cell(0, 5.5, stripped, align="J")
            pdf.ln(1)

    return bytes(pdf.output())
